import pandas as pd
import os
from datetime import datetime
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    Document = None

class ReportGenerator:
    def __init__(self, inventory_df):
        self.df = inventory_df.copy()

    def apply_filters(self, conditions):
        """
        conditions: list of dicts {'logic': str, 'field': str, 'operator': str, 'value': str}
        operators: 'Equals', 'Contains', '>', '<', '>=', '<=', 'Not Equals', 'Modulo', 'Above', 'Below'
        logic: 'START', 'AND', 'OR', 'AND NOT', 'OR NOT', 'XOR'
        """
        if self.df.empty:
            return pd.DataFrame()
            
        if not conditions:
            return self.df

        # Initialize the master mask
        # If the first condition is valid, we'll set it.
        # Otherwise, we start with all True.
        final_mask = pd.Series([True] * len(self.df), index=self.df.index)
        first_valid = True

        for i, cond in enumerate(conditions):
            logic = cond.get('logic', 'AND')
            if i == 0: logic = 'START' # Force start logic for first item
            
            field = cond['field']
            op = cond['operator']
            val = cond['value']

            if field not in self.df.columns:
                continue
            
            # Handle empty values
            if not val and op not in ['Is Empty', 'Is Not Empty']:
                continue

            try:
                current_mask = self._build_mask(field, op, val)
                
                if logic == 'START' or first_valid:
                    final_mask = current_mask
                    first_valid = False
                elif logic == 'AND':
                    final_mask = final_mask & current_mask
                elif logic == 'OR':
                    final_mask = final_mask | current_mask
                elif logic == 'AND NOT':
                    final_mask = final_mask & (~current_mask)
                elif logic == 'OR NOT':
                    final_mask = final_mask | (~current_mask)
                elif logic == 'XOR':
                    final_mask = final_mask ^ current_mask
                else:
                    # Default AND
                    final_mask = final_mask & current_mask

            except Exception as e:
                print(f"Error applying filter {cond}: {e}")
                # Don't break the whole chain?
                pass

        return self.df[final_mask]

    def _build_mask(self, field, op, val):
        """Generates a boolean Series mask for a single condition."""
        # Helper for type conversion
        df = self.df
        col_type = df[field].dtype
        mask = pd.Series([False] * len(df), index=df.index)
        
        try:
            if op == 'Equals':
                if col_type == 'object':
                     mask = df[field].astype(str).str.lower() == str(val).lower()
                else:
                    mask = df[field] == val
            
            elif op == 'Contains':
                mask = df[field].astype(str).str.contains(val, case=False, na=False)
            
            elif op == 'Not Equals':
                if col_type == 'object':
                     mask = df[field].astype(str).str.lower() != str(val).lower()
                else:
                    mask = df[field] != val

            elif op == 'Modulo':
                if "=" in str(val):
                    div, rem = map(int, str(val).split('='))
                    nums = pd.to_numeric(df[field], errors='coerce').fillna(0)
                    mask = (nums % div == rem)

            elif op in ['>', '<', '>=', '<=', 'Above', 'Below']:
                # Handle Date/Numeric
                is_date = pd.api.types.is_datetime64_any_dtype(df[field])
                compare_val = val
                
                if is_date:
                    try:
                        compare_val = pd.to_datetime(val)
                    except: pass
                else:
                    try:
                        compare_val = float(val)
                    except: pass
                
                if op == '>' or op == 'Above':
                    mask = df[field] > compare_val
                elif op == '<' or op == 'Below':
                    mask = df[field] < compare_val
                elif op == '>=':
                    mask = df[field] >= compare_val
                elif op == '<=':
                    mask = df[field] <= compare_val
            
            elif op == 'Is Empty':
                 mask = df[field].isna() | (df[field] == '')
                 
            elif op == 'Is Not Empty':
                 mask = df[field].notna() & (df[field] != '')
        except Exception as e:
            print(f"Mask build error: {e}")
            return mask # False mask on error
            
        return mask

    def apply_limit(self, df, limit):
        try:
            lim = int(limit)
            if lim > 0:
                return df.head(lim)
        except:
            pass
        return df

    def apply_custom_expression(self, df, expression):
        if not expression or not str(expression).strip():
            return df
        try:
            # Use pandas query engine
            # Replace 'unique_id' with backticks if needed, but pandas query handles it.
            # Security: query() uses eval() under the hood but restricted to DF context.
            return df.query(str(expression))
        except Exception as e:
            print(f"Custom Query Error: {e}")
            return pd.DataFrame()


    def export(self, data, columns, format_type, filepath, include_serial=False):
        """
        data: Filtered DataFrame
        columns: List of columns to include
        format_type: 'excel', 'pdf', 'word'
        """
        if data.empty:
            return False, "No data to export"
            
        # Select only requested columns
        # Respect user order (columns list order)
        valid_cols = [c for c in columns if c in data.columns]
        export_data = data[valid_cols].copy()

        if include_serial:
            export_data.insert(0, 'S.No', range(1, 1 + len(export_data)))

        try:
            if format_type == 'excel':
                export_data.to_excel(filepath, index=False)
            
            elif format_type == 'pdf':
                self._export_pdf(export_data, filepath)
                
            elif format_type == 'word':
                if Document is None:
                    return False, "python-docx library not installed."
                self._export_word(export_data, filepath)
            
            return True, f"Successfully exported to {filepath}"
            
        except Exception as e:
            return False, str(e)

    def _export_pdf(self, df, filepath):
        doc = SimpleDocTemplate(filepath, pagesize=landscape(A4))
        elements = []
        styles = getSampleStyleSheet()
        
        # Title
        title = Paragraph(f"Stock Report - {datetime.now().strftime('%Y-%m-%d')}", styles['Title'])
        elements.append(title)
        elements.append(Spacer(1, 20))

        # Data conversion
        # Headers
        data = [df.columns.tolist()]
        # Rows
        for index, row in df.iterrows():
            data.append([str(x) for x in row.tolist()])

        # Table
        t = Table(data)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 0), (-1, -1), 8),  # Small font to fit many cols
        ]))
        
        elements.append(t)
        doc.build(elements)

    def _export_word(self, df, filepath):
        doc = Document()
        doc.add_heading(f"Stock Report - {datetime.now().strftime('%Y-%m-%d')}", 0)

        # Add table
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'

        # Header
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = str(col_name)

        # Rows
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)

        doc.save(filepath)